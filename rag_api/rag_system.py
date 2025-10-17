import os
import uuid
import io
from typing import List, Dict, Any, Annotated, Sequence, Literal, Optional, Union, Tuple
from decouple import config
from langchain_core.documents import Document
from langchain_core.messages import BaseMessage, HumanMessage, AIMessage, SystemMessage
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langgraph.graph import StateGraph, END, START
from langgraph.graph.message import add_messages
from typing_extensions import TypedDict
from operator import add
import json

# For handling PDF and DOCX files
try:
    import docx
    from PyPDF2 import PdfReader
    PDF_DOCX_SUPPORT = True
except ImportError:
    PDF_DOCX_SUPPORT = False
    print("⚠️ PDF and DOCX support not available. Install dependencies with: pip install python-docx PyPDF2")


class RAGState(TypedDict):
    """State for the RAG System"""
    messages: Annotated[Sequence[BaseMessage], add_messages]
    context: str
    sources: List[str]
    session_id: str
    instructions: Optional[str]


class RAGTool(TypedDict):
    """Tool definition for the RAG Agent"""
    name: str
    description: str
    args_schema: Dict


class RAGSystem:
    def __init__(self):
        self.llm = None
        self.llm_client = None
        self.embed_client = None
        self.connection_string = None
        self.setup_llm()
        self.setup_database()
    
    def setup_llm(self):
        """Initialize LLM and embedding models"""
        print("🔄 Initializing RAG System...")
        
        # Debug environment variables
        print(f"🔍 AZURE_API_KEY: {'*' * 10 if config('AZURE_API_KEY', default='') else 'NOT SET'}")
        print(f"🔍 AZURE_ENDPOINT: {config('AZURE_ENDPOINT', default='NOT SET')}")
        print(f"🔍 EMBED_API_KEY: {'*' * 10 if config('EMBED_API_KEY', default='') else 'NOT SET'}")
        print(f"🔍 EMBED_ENDPOINT: {config('EMBED_ENDPOINT', default='NOT SET')}")
        
        # Initialize clients step by step with detailed error reporting
        self.llm = None
        self.llm_client = None
        self.embed_client = None
        
        try:
            from openai import AzureOpenAI
            print("✅ OpenAI module imported successfully")
            
            # Test LLM client initialization
            print(f"🔄 Initializing LLM client with endpoint: {config('AZURE_ENDPOINT')}")
            self.llm_client = AzureOpenAI(
                api_key=config("AZURE_API_KEY"),
                api_version=config("AZURE_API_VERSION"),
                azure_endpoint=config("AZURE_ENDPOINT")
            )
            self.llm_deployment = config("AZURE_DEPLOYMENT")
            print("✅ LLM client initialized successfully")
            
        except Exception as e:
            print(f"❌ Error initializing LLM client: {e}")
            self.llm_client = None
        
        try:
            from openai import AzureOpenAI
            
            # Test embedding client initialization
            print(f"🔄 Initializing embedding client with endpoint: {config('EMBED_ENDPOINT')}")
            self.embed_client = AzureOpenAI(
                api_key=config("EMBED_API_KEY"),
                api_version=config("EMBED_API_VERSION"),
                azure_endpoint=config("EMBED_ENDPOINT")
            )
            self.embed_deployment = config("EMBED_DEPLOYMENT")
            print("✅ Embedding client initialized successfully")
            
        except Exception as e:
            print(f"❌ Error initializing embedding client: {e}")
            self.embed_client = None
        
        # Report final status
        if self.llm_client and self.embed_client:
            print("✅ RAG System fully initialized with direct Azure OpenAI clients")
        elif self.llm_client:
            print("⚠️ RAG System partially initialized (LLM only, no embeddings)")
        elif self.embed_client:
            print("⚠️ RAG System partially initialized (embeddings only, no LLM)")
        else:
            print("❌ RAG System initialization failed completely")
    
    def generate_embedding(self, text: str) -> List[float]:
        """Generate embedding using direct Azure OpenAI client"""
        if not self.embed_client:
            print("❌ Embedding client not available")
            return []
            
        if not text or not text.strip():
            print("❌ Cannot generate embedding for empty text")
            return []
        
        try:
            # Truncate text if it's too long (embedding models often have token limits)
            # Ada embedding model has a limit of 8191 tokens, which is roughly 32k characters
            max_chars = 32000
            if len(text) > max_chars:
                print(f"⚠️ Truncating text from {len(text)} to {max_chars} characters for embedding")
                text = text[:max_chars]
                
            response = self.embed_client.embeddings.create(
                input=text,
                model=self.embed_deployment
            )
            
            if not response or not response.data or len(response.data) == 0:
                print("❌ Empty embedding response")
                return []
                
            return response.data[0].embedding
        except Exception as e:
            print(f"❌ Error generating embedding: {str(e)}")
            import traceback
            traceback.print_exc()
            return []
    
    def similarity_search(self, vector_store, query: str, k: int = 3) -> List[Document]:
        """Perform similarity search in development mode"""
        if config('DEBUG', default=True, cast=bool):
            # Development mode with cosine similarity
            # If query is empty, just return the first k documents without calculating similarity
            if not query.strip():
                print("Empty query, returning first documents")
                documents = []
                for item in vector_store[:k]:
                    if 'document' in item:
                        documents.append(item['document'])
                    elif 'content' in item:
                        # Create a Document object if there's content
                        from langchain.schema import Document as LangchainDocument
                        documents.append(LangchainDocument(page_content=item['content']))
                return documents
                
            # For non-empty queries, calculate similarity
            query_embedding = self.generate_embedding(query)
            if not query_embedding:
                print("Could not generate embedding for query")
                return []
            
            # Calculate similarities
            similarities = []
            for item in vector_store:
                if 'embedding' in item and item['embedding']:
                    similarity = self.cosine_similarity(query_embedding, item['embedding'])
                    similarities.append((similarity, item['document']))
            
            # Sort by similarity and return top k
            similarities.sort(key=lambda x: x[0], reverse=True)
            return [doc for _, doc in similarities[:k]]
        else:
            # Production mode with PGVector
            if not query.strip():
                # Handle empty query in production mode
                try:
                    return vector_store.max_marginal_relevance_search("", k=k)
                except:
                    # Fallback to regular search
                    return vector_store.similarity_search("document", k=k)
            else:
                return vector_store.similarity_search(query, k=k)
    
    def cosine_similarity(self, vec1: List[float], vec2: List[float]) -> float:
        """Calculate cosine similarity between two vectors"""
        try:
            import math
            dot_product = sum(a * b for a, b in zip(vec1, vec2))
            magnitude1 = math.sqrt(sum(a * a for a in vec1))
            magnitude2 = math.sqrt(sum(a * a for a in vec2))
            if magnitude1 == 0 or magnitude2 == 0:
                return 0
            return dot_product / (magnitude1 * magnitude2)
        except Exception as e:
            print(f"Error calculating similarity: {e}")
            return 0
    
    def setup_database(self):
        """Setup database connection"""
        try:
            if config('DEBUG', default=True, cast=bool):
                # For development, we'll use in-memory storage
                self.connection_string = None
                print("✅ Using in-memory storage for development")
            else:
                # For production, use PostgreSQL
                db_name = config("ALLOYDB_DATABASE")
                self.connection_string = f"postgresql://{config('ALLOYDB_USERNAME')}:{config('ALLOYDB_PASS')}@{config('ALLOYDB_SERVER')}:5432/{db_name}"
                print(f"✅ Using PostgreSQL: {config('ALLOYDB_SERVER')}")
        except Exception as e:
            print(f"❌ Database setup error: {e}")
            self.connection_string = None
    
    def create_vector_store(self, session_id: str):
        """Create a new vector store for a session"""
        collection_name = f"rag_collection_{str(session_id).replace('-', '_')}"
        print(f"Creating/retrieving vector store for collection: {collection_name}")
        
        # For development with SQLite, use in-memory storage
        if config('DEBUG', default=True, cast=bool):
            # Use a simple in-memory storage for development
            if not hasattr(self, '_dev_storage'):
                print("Initializing in-memory vector store")
                self._dev_storage = {}
            
            if collection_name not in self._dev_storage:
                print(f"Creating new collection: {collection_name}")
                self._dev_storage[collection_name] = []
            else:
                print(f"Using existing collection with {len(self._dev_storage[collection_name])} documents")
                
            return self._dev_storage[collection_name]
        else:
            # Production mode disabled for now - use development storage
            if not hasattr(self, '_dev_storage'):
                print("Initializing in-memory vector store for production mode")
                self._dev_storage = {}
            
            if collection_name not in self._dev_storage:
                print(f"Creating new collection in production mode: {collection_name}")
                self._dev_storage[collection_name] = []
            
            return self._dev_storage[collection_name]
    
    def extract_text_from_file(self, file_content, file_type):
        """Extract text content from different file types"""
        if file_type == 'txt':
            # For text files, just return the content
            return file_content
        
        elif file_type == 'pdf' and PDF_DOCX_SUPPORT:
            # For PDF files
            try:
                pdf_file = io.BytesIO(file_content)
                pdf_reader = PdfReader(pdf_file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
            except Exception as e:
                print(f"❌ Error extracting text from PDF: {e}")
                return ""
        
        elif file_type == 'docx' and PDF_DOCX_SUPPORT:
            # For DOCX files
            try:
                docx_file = io.BytesIO(file_content)
                doc = docx.Document(docx_file)
                text = ""
                for para in doc.paragraphs:
                    text += para.text + "\n"
                return text
            except Exception as e:
                print(f"❌ Error extracting text from DOCX: {e}")
                return ""
        
        return ""
    
    def generate_qa_for_chunk(self, chunk_content: str) -> Tuple[str, str, str]:
        """Generate question-answer pair and summary for a chunk"""
        if not self.llm_client:
            # Return empty values if LLM is not available
            return "", "", ""
        
        try:
            response = self.llm_client.chat.completions.create(
                model=self.llm_deployment,
                messages=[
                    {"role": "system", "content": "You are an expert at extracting meaningful information from text and creating high-quality question-answer pairs. Generate a concise summary, a specific question, and a comprehensive answer for the given text chunk."},
                    {"role": "user", "content": f"""For the following text chunk, please provide:
1. A concise summary (2-3 sentences)
2. A specific and meaningful question that could be asked about this content
3. A comprehensive answer to that question

Text chunk:
{chunk_content}

Format your response as a JSON object with these keys: "summary", "question", "answer"
"""}
                ],
                temperature=0.3,
                response_format={"type": "json_object"}
            )
            
            try:
                result = json.loads(response.choices[0].message.content)
                return result.get("summary", ""), result.get("question", ""), result.get("answer", "")
            except Exception as e:
                print(f"Error parsing QA generation response: {e}")
                return "", "", ""
                
        except Exception as e:
            print(f"Error generating QA pair: {e}")
            return "", "", ""
    
    def process_documents(self, documents_content: List[str], session_id: str, document_objects=None, file_types=None) -> Tuple[bool, List[Dict]]:
        """Process documents and create vector index with enhanced metadata"""
        if not self.embed_client:
            print("❌ Embedding client not available, skipping vector processing")
            return False, []
            
        try:
            print(f"Processing {len(documents_content)} documents for session {session_id}")
            
            # Create vector store for this session
            vector_store = self.create_vector_store(session_id)
            print(f"Vector store created/retrieved: {type(vector_store)}")
            
            # Text splitter for chunking
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
            )
            
            # Store document details for return
            all_chunks_metadata = []
            
            # Process each document
            for i, content in enumerate(documents_content):
                if not content or not content.strip():
                    print(f"Skipping empty document at index {i}")
                    continue
                    
                print(f"Processing document {i} with {len(content)} characters")
                doc_chunks = text_splitter.split_text(content)
                print(f"Document {i} split into {len(doc_chunks)} chunks")
                document_chunks_metadata = []
                
                for j, chunk in enumerate(doc_chunks):
                    try:
                        # Generate Q&A and summary for the chunk
                        print(f"Generating Q&A for document {i}, chunk {j}")
                        summary, question, answer = self.generate_qa_for_chunk(chunk)
                        
                        # Create enhanced metadata
                        chunk_metadata = {
                            "source": f"document_{i}" if document_objects is None else str(document_objects[i].id),
                            "chunk": j,
                            "summary": summary,
                            "question": question,
                            "answer": answer
                        }
                        
                        # Create document
                        from langchain.schema import Document as LangchainDocument
                        doc = LangchainDocument(
                            page_content=chunk,
                            metadata=chunk_metadata
                        )
                        
                        # Generate embedding
                        print(f"Generating embedding for document {i}, chunk {j}")
                        embedding = self.generate_embedding(doc.page_content)
                        
                        if embedding:
                            # Store document with embedding
                            doc_with_embedding = {
                                'document': doc,
                                'embedding': embedding,
                                'content': doc.page_content,
                                'metadata': chunk_metadata
                            }
                            
                            # Add to vector store
                            if config('DEBUG', default=True, cast=bool) and isinstance(vector_store, list):
                                vector_store.append(doc_with_embedding)
                                print(f"Added chunk to in-memory vector store (total: {len(vector_store)})")
                            else:
                                # Handle other vector store implementations
                                try:
                                    vector_store.add_documents([doc])
                                    print(f"Added chunk to vector database")
                                except Exception as e:
                                    print(f"Error adding to vector store: {str(e)}")
                                    # Still try the list approach as fallback
                                    vector_store.append(doc_with_embedding)
                        else:
                            print(f"Warning: Could not generate embedding for chunk {j} of document {i}")
                    except Exception as e:
                        print(f"Error processing chunk {j} of document {i}: {str(e)}")
                        # Continue with next chunk
                    
                    # Store metadata for return
                    document_chunks_metadata.append({
                        'content': chunk,
                        'summary': summary,
                        'question': question,
                        'answer': answer,
                        'chunk_index': j,
                        'document_index': i
                    })
                
                all_chunks_metadata.extend(document_chunks_metadata)
            
            print(f"✅ Stored {len(all_chunks_metadata)} document chunks with embeddings for session {session_id}")
            return True, all_chunks_metadata
            
        except Exception as e:
            print(f"❌ Error processing documents: {e}")
            return False, []
    
    def get_existing_vector_store(self, session_id: str):
        """Get existing vector store for a session"""
        return self.create_vector_store(session_id)
    
    def generate_title_from_content(self, content: str) -> str:
        """Generate a title from document content"""
        if not self.llm_client:
            # Generate a simple title from content
            words = content.split()[:5]
            return " ".join(words) + "..." if len(words) == 5 else " ".join(words)
            
        # Truncate content to first 1000 characters to avoid token limits
        truncated_content = content[:1000]
        
        try:
            response = self.llm_client.chat.completions.create(
                model=self.llm_deployment,
                messages=[
                    {"role": "system", "content": "Generate a short, descriptive title (maximum 10 words) that summarizes what this document is about. Return only the title, nothing else."},
                    {"role": "user", "content": f"Content: {truncated_content}"}
                ],
                temperature=0.3,
                max_tokens=50
            )
            
            title = response.choices[0].message.content.strip()
            title = title.replace("Title:", "").strip()
            return title[:100]  # Limit to 100 characters
        except Exception as e:
            print(f"Error generating title: {e}")
            # Generate a simple title from content
            words = content.split()[:5]
            return " ".join(words) + "..." if len(words) == 5 else " ".join(words)
    
    def generate_suggestive_questions(self, session_id: str) -> List[str]:
        """Generate suggestive questions based on the documents"""
        try:
            vector_store = self.get_existing_vector_store(session_id)
            
            # Print debug information
            print(f"Generating suggestive questions for session: {session_id}")
            print(f"Vector store: {vector_store}")
            
            # Get documents from the vector store
            if config('DEBUG', default=True, cast=bool):
                # In development mode, vector_store is a list of dictionaries
                if not vector_store or len(vector_store) == 0:
                    print("Vector store is empty, using fallback questions")
                    return self._get_fallback_questions()
                
                # Get text content from the first 3 documents or all if less than 3
                sample_content = ""
                for i, item in enumerate(vector_store[:3]):
                    if 'document' in item and hasattr(item['document'], 'page_content'):
                        sample_content += item['document'].page_content + "\n\n"
                    elif 'content' in item:
                        sample_content += item['content'] + "\n\n"
            else:
                # In production mode, use similarity search
                sample_docs = self.similarity_search(vector_store, "", k=3)
                if not sample_docs:
                    print("No documents returned from similarity search, using fallback questions")
                    return self._get_fallback_questions()
                    
                sample_content = "\n".join([doc.page_content for doc in sample_docs if hasattr(doc, 'page_content')])
            
            # Check if we have content to generate questions from
            if not sample_content.strip():
                print("No content available to generate questions, using fallback questions")
                return self._get_fallback_questions()
                
            print(f"Sample content length: {len(sample_content)}")
            
            # Generate questions using the LLM
            if self.llm_client:
                print("Calling LLM to generate questions")
                response = self.llm_client.chat.completions.create(
                    model=self.llm_deployment,
                    messages=[
                        {"role": "system", "content": "Generate 6 diverse and insightful questions that users might ask about the document content. Questions should cover different aspects, be specific and actionable, and range from basic to analytical. Return only the questions, one per line, numbered 1-6."},
                        {"role": "user", "content": f"Content: {sample_content[:2000]}"}
                    ],
                    temperature=0.7,
                    max_tokens=300
                )
                response_text = response.choices[0].message.content
                print(f"LLM response: {response_text}")
            else:
                print("LLM client not available, using fallback questions")
                return self._get_fallback_questions()
            
            # Parse questions
            questions = []
            for line in response_text.split('\n'):
                line = line.strip()
                if line and (line[0].isdigit() or line.startswith('-') or line.startswith('•')):
                    # Remove numbering and clean up
                    question = line
                    if '.' in question:
                        question = question.split('.', 1)[-1].strip()
                    # Ensure question ends with question mark
                    if not question.endswith('?'):
                        question = question.rstrip('?') + '?'
                    if len(question) > 10:  # Ensure it's a real question
                        questions.append(question)
            
            print(f"Parsed questions: {questions}")
            
            # Fallback questions if generation fails
            if len(questions) < 6:
                print(f"Only found {len(questions)} questions, adding fallbacks")
                questions.extend(self._get_fallback_questions())
            
            # Return up to 6 questions
            return questions[:6]
            
        except Exception as e:
            print(f"Error generating questions: {str(e)}")
            import traceback
            traceback.print_exc()
            # Return default questions
            return self._get_fallback_questions()
            
    def _get_fallback_questions(self) -> List[str]:
        """Return a set of fallback questions"""
        return [
            "What are the main topics covered in these documents?",
            "Can you summarize the key points?",
            "What are the most important findings or conclusions?",
            "Are there any specific recommendations mentioned?",
            "What problems or challenges are discussed?",
            "How can this information be applied practically?"
        ]
    
    # Methods for React agent pattern have been removed
    
    def create_rag_graph(self) -> StateGraph:
        """Create LangGraph workflow for simple RAG system"""
        
        def retrieve_context(state: RAGState) -> RAGState:
            """Retrieve relevant context from vector store"""
            try:
                session_id = state["session_id"]
                # Get the last user message
                last_message = state["messages"][-1]
                question = last_message.content
                
                vector_store = self.get_existing_vector_store(session_id)
                relevant_docs = self.similarity_search(vector_store, question, k=3)
                
                context = "\n\n".join([doc.page_content for doc in relevant_docs])
                sources = [doc.metadata.get("source", f"Document {i+1}") for i, doc in enumerate(relevant_docs)]
                
                state["context"] = context
                state["sources"] = sources
                
            except Exception as e:
                print(f"Error retrieving context: {e}")
                state["context"] = "No relevant context found."
                state["sources"] = []
            
            return state
        
        def generate_answer(state: RAGState) -> RAGState:
            """Generate an answer using the context"""
            # If LLM client is not available, return an error
            if not self.llm_client:
                error_response = "I apologize, but the AI system is not properly configured. Please check the Azure OpenAI settings."
                state["messages"].append(AIMessage(content=error_response))
                return state
            
            last_message = state["messages"][-1]
            context = state["context"]
            
            # Get custom instructions if available
            instructions = state.get("instructions", "")
            
            # Create system prompt with context
            system_prompt = f"""You are a helpful AI assistant. Answer the user's questions based on the provided context.
            
            CONTEXT:
            {context}
            
            INSTRUCTIONS:
            {instructions}
            
            - Respond directly to the user's question
            - If the context doesn't contain the information needed, say so
            - Be concise and helpful
            - Do not mention that you were given a context"""
            
            # Build conversation history for OpenAI format
            openai_messages = [
                {"role": "system", "content": system_prompt}
            ]
            
            # Add conversation history
            for msg in state["messages"]:
                if isinstance(msg, HumanMessage):
                    openai_messages.append({"role": "user", "content": msg.content})
                elif isinstance(msg, AIMessage):
                    openai_messages.append({"role": "assistant", "content": msg.content})
                elif isinstance(msg, SystemMessage):
                    openai_messages.append({"role": "system", "content": msg.content})
            
            try:
                # Call OpenAI API
                response = self.llm_client.chat.completions.create(
                    model=self.llm_deployment,
                    messages=openai_messages,
                    temperature=0.1,
                    max_tokens=1000
                )
                
                ai_response = response.choices[0].message.content
                state["messages"].append(AIMessage(content=ai_response))
                
            except Exception as e:
                print(f"Error generating answer: {e}")
                error_response = "I apologize, but I encountered an error while processing your request. Please try again."
                state["messages"].append(AIMessage(content=error_response))
            
            return state
        
        # Create the graph
        graph = StateGraph(RAGState)
        
        # Add nodes
        graph.add_node("retrieve", retrieve_context)
        graph.add_node("generate", generate_answer)
        
        # Add edges
        graph.add_edge(START, "retrieve")
        graph.add_edge("retrieve", "generate")
        graph.add_edge("generate", END)
        
        # Compile the graph
        return graph.compile()
    
    def chat(self, session_id: str, question: str, chat_history: List[Dict] = None, instructions: str = None, stream: bool = False):
        """Process a chat message through the simple RAG system"""
        # Check if system is properly initialized
        if not self.llm_client:
            response = {
                "answer": "I apologize, but the AI system is not properly configured. Please check the Azure OpenAI settings and ensure all required environment variables are set.",
                "sources": [],
                "context": "System not initialized"
            }
            if stream:
                yield response
                return
            else:
                return response
        
        # Prepare messages
        messages = []
        if chat_history:
            for msg in chat_history:
                if msg["message_type"] == "user":
                    messages.append(HumanMessage(content=msg["content"]))
                elif msg["message_type"] == "assistant":
                    messages.append(AIMessage(content=msg["content"]))
        
        # Add current question
        messages.append(HumanMessage(content=question))
        
        # Format instructions if provided
        formatted_instructions = ""
        if instructions:
            formatted_instructions = f"""
            {instructions}
            
            When responding to the user, always follow these guidelines:
            - Use the information from the documents to provide accurate answers
            - When information is available in Q&A pairs, prioritize using those
            - Use summaries to give concise overviews when appropriate
            - If the information isn't in the knowledge base, clearly state that
            """
        
        # For streaming, we need to handle it differently
        if stream:
            # First, retrieve context
            context = ""
            sources = []
            
            try:
                vector_store = self.get_existing_vector_store(session_id)
                relevant_docs = self.similarity_search(vector_store, question, k=3)
                
                context = "\n\n".join([doc.page_content for doc in relevant_docs])
                sources = [doc.metadata.get("source", f"Document {i+1}") for i, doc in enumerate(relevant_docs)]
            except Exception as e:
                print(f"Error retrieving context: {str(e)}")
                context = "No relevant context found."
            
            # Create system prompt with context
            system_prompt = f"""You are a helpful AI assistant. Answer the user's questions based on the provided context.
            
            CONTEXT:
            {context}
            
            INSTRUCTIONS:
            {formatted_instructions}
            
            - Respond directly to the user's question
            - If the context doesn't contain the information needed, say so
            - Be concise and helpful
            - Do not mention that you were given a context"""
            
            # Build conversation history for OpenAI format
            openai_messages = [
                {"role": "system", "content": system_prompt}
            ]
            
            # Add conversation history
            for msg in messages:
                if isinstance(msg, HumanMessage):
                    openai_messages.append({"role": "user", "content": msg.content})
                elif isinstance(msg, AIMessage):
                    openai_messages.append({"role": "assistant", "content": msg.content})
                elif isinstance(msg, SystemMessage):
                    openai_messages.append({"role": "system", "content": msg.content})
            
            try:
                print(f"Starting streaming response for session {session_id}")
                print(f"Number of messages being sent: {len(openai_messages)}")
                print(f"Context length: {len(context)}")
                
                # Verify LLM client
                if not self.llm_client:
                    print("LLM client is None!")
                    raise Exception("LLM client not initialized")
                
                # Check if deployment is valid
                if not self.llm_deployment:
                    print("LLM deployment is None!")
                    raise Exception("LLM deployment not specified")
                    
                print(f"Using LLM deployment: {self.llm_deployment}")
                
                # Stream OpenAI API response
                try:
                    # Print messages length to debug potential token limit issues
                    message_lengths = [len(msg["content"]) for msg in openai_messages]
                    total_message_length = sum(message_lengths)
                    print(f"Total message length: {total_message_length} characters")
                    print(f"System message length: {len(openai_messages[0]['content'])} characters")
                    print(f"Number of messages: {len(openai_messages)}")
                    
                    response_stream = self.llm_client.chat.completions.create(
                        model=self.llm_deployment,
                        messages=openai_messages,
                        temperature=0.1,
                        max_tokens=1000,
                        stream=True
                    )
                    
                    print("Successfully initiated streaming response")
                except Exception as api_error:
                    print(f"API ERROR: Failed to create completion stream: {str(api_error)}")
                    import traceback
                    traceback.print_exc()
                    # Yield an error response that can be handled by the caller
                    yield {
                        "chunk": f"Error initiating streaming response: {str(api_error)}",
                        "full_response": f"Error initiating streaming response: {str(api_error)}",
                        "sources": sources,
                        "context": context,
                        "finish_reason": "api_error",
                        "done": True
                    }
                    return  # Exit the generator
                
                # Stream chunks back to the caller
                full_response = ""
                chunk_count = 0
                
                try:
                    for chunk in response_stream:
                        chunk_count += 1
                        
                        # Debug log to see the actual structure of each chunk
                        if chunk_count <= 3:
                            print(f"DEBUG - Chunk {chunk_count} structure: {chunk}")
                        
                        # Safely access the choices and delta
                        if hasattr(chunk, 'choices') and chunk.choices and len(chunk.choices) > 0:
                            choice = chunk.choices[0]
                            if hasattr(choice, 'delta') and hasattr(choice.delta, 'content') and choice.delta.content:
                                content = choice.delta.content
                                full_response += content
                                
                                # Get finish reason safely
                                finish_reason = None
                                if hasattr(choice, 'finish_reason'):
                                    finish_reason = choice.finish_reason
                                
                                # Yield the chunk
                                response_chunk = {
                                    "chunk": content,
                                    "full_response": full_response,
                                    "sources": sources,
                                    "context": context,
                                    "finish_reason": finish_reason
                                }
                                
                                # Log every 10th chunk to avoid excessive logging
                                if chunk_count % 10 == 0:
                                    print(f"Yielding chunk #{chunk_count}, response length: {len(full_response)}")
                                    
                                yield response_chunk
                        else:
                            # Log chunks that don't have the expected structure
                            print(f"DEBUG - Received chunk without expected structure: {chunk}")
                except Exception as stream_error:
                    print(f"STREAM ERROR during streaming: {str(stream_error)}")
                    import traceback
                    traceback.print_exc()
                    # Yield an error response that can be handled by the caller
                    yield {
                        "chunk": f"Error during streaming: {str(stream_error)}",
                        "full_response": full_response if full_response else f"Error during streaming: {str(stream_error)}",
                        "sources": sources,
                        "context": context,
                        "finish_reason": "stream_error",
                        "done": True
                    }
                
                print(f"Streaming complete, yielded {chunk_count} chunks, total response length: {len(full_response) if full_response else 0}")
                
                # Only send a final chunk if we haven't yielded an error already
                if full_response or not chunk_count:
                    # Send a final chunk with the complete response
                    final_response = full_response if full_response else "I apologize, but I couldn't generate a proper response. Please try again."
                    print(f"Sending final chunk with response length: {len(final_response)}")
                    yield {
                        "chunk": "",
                        "full_response": final_response,
                        "sources": sources,
                        "context": context,
                        "finish_reason": "stop",
                        "done": True
                    }
                
            except Exception as e:
                print(f"Error generating streaming answer: {str(e)}")
                error_response = "I apologize, but I encountered an error while processing your request. Please try again."
                yield {
                    "chunk": error_response,
                    "full_response": error_response,
                    "sources": sources,
                    "context": context,
                    "finish_reason": "error",
                    "done": True
                }
        else:
            # Non-streaming mode - use the graph
            graph = self.create_rag_graph()
            
            # Create initial state
            initial_state = {
                "messages": messages,
                "context": "",
                "sources": [],
                "session_id": session_id,
                "instructions": formatted_instructions
            }
            
            result = graph.invoke(initial_state)
            
            # Get the last AI message
            ai_response = result["messages"][-1].content
            
            return {
                "answer": ai_response,
                "sources": result["sources"],
                "context": result["context"]
            }
    
    def add_documents_to_existing_session(self, session_id: str, documents_content: List[str], document_objects=None, file_types=None):
        """Add new documents to an existing session's vector store with enhanced metadata"""
        try:
            # Delegate to the process_documents method for consistency
            return self.process_documents(
                documents_content, 
                session_id, 
                document_objects,
                file_types
            )
        except Exception as e:
            print(f"Error adding documents to session: {e}")
            return False, []
    
    def delete_session_vector_store(self, session_id: str):
        """Delete the vector store collection for a session"""
        try:
            collection_name = f"rag_collection_{str(session_id).replace('-', '_')}"
            
            if config('DEBUG', default=True, cast=bool):
                # Development mode - remove from in-memory storage
                if hasattr(self, '_dev_storage') and collection_name in self._dev_storage:
                    del self._dev_storage[collection_name]
                return True
            else:
                # Production mode disabled for now
                return True
                
        except Exception as e:
            print(f"Error deleting vector store: {e}")
            return False


# Global RAG system instance - will be initialized on first use
rag_system = None

def get_rag_system():
    global rag_system
    if rag_system is None:
        try:
            rag_system = RAGSystem()
            print("✅ RAG System initialized successfully")
        except Exception as e:
            print(f"❌ Failed to initialize RAG System: {e}")
            # Return None to indicate failure - views should handle this
            raise Exception(f"Failed to initialize RAG system: {e}")
    return rag_system